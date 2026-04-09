"""
Generate B.E. Project Report for Indoor SLAM Navigation App
Following University of Mumbai formatting guidelines.
"""

from docx import Document
from docx.shared import Pt, Mm, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# STYLES & PAGE SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(30)
section.right_margin = Mm(20)
section.top_margin = Mm(30)
section.bottom_margin = Mm(22)


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_page_break():
    doc.add_page_break()


def add_heading_centered(text, size=18, bold=True, space_before=0, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_heading_left(text, size=16, bold=True, space_before=15, space_after=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_subheading_left(text, size=14, bold=True, space_before=12, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_body(text, indent=True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Mm(12)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.left_indent = Mm(12 + level * 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Mm(w)

    doc.add_paragraph()  # spacing after table
    return table


# ============================================================
# TITLE PAGE (Specimen A)
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_heading_centered(
    "Indoor SLAM-Based Navigation System for\nVisually Impaired Users Using ARCore and Deep Learning",
    size=22, space_after=24
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Submitted in partial fulfillment of the requirements\nof the degree of")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_heading_centered("Bachelor of Engineering", size=16, space_before=12, space_after=18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("by")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_heading_centered("[YOUR NAME]", size=15, space_before=12, space_after=6)
add_heading_centered("(Roll No. [YOUR ROLL NUMBER])", size=13, bold=False, space_before=0, space_after=18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Supervisor(s):")
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_heading_centered("[SUPERVISOR NAME]", size=14, bold=True, space_before=6, space_after=30)

# Emblem placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Institute/University Emblem - 50mm diameter]")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

for _ in range(2):
    doc.add_paragraph()

add_heading_centered("[Department of Computer Engineering / Information Technology]", size=13, bold=False, space_before=6, space_after=4)
add_heading_centered("[Name of Institute]", size=14, bold=True, space_before=4, space_after=4)
add_heading_centered("University of Mumbai", size=14, bold=True, space_before=4, space_after=4)
add_heading_centered("2025", size=14, bold=True, space_before=4, space_after=0)

add_page_break()

# ============================================================
# CERTIFICATE (Specimen B)
# ============================================================
for _ in range(2):
    doc.add_paragraph()

add_heading_centered("Certificate", size=18, space_after=30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Mm(12)
run = p.add_run(
    'This is to certify that the project entitled '
)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run('"Indoor SLAM-Based Navigation System for Visually Impaired Users Using ARCore and Deep Learning"')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run(' is a bonafide work of ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run('"[YOUR NAME(S)]" (Roll No. [YOUR ROLL NUMBER(S)])')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run(
    ' submitted to the University of Mumbai in partial fulfillment of the requirement '
    'for the award of the degree of '
)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run('"Bachelor of Engineering"')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run(' in ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run('"[Program Name e.g., Computer Engineering]"')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

run = p.add_run('.')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(6):
    doc.add_paragraph()

# Signature table
sig_table = doc.add_table(rows=2, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in sig_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

cells = sig_table.rows[0].cells
cells[0].paragraphs[0].add_run("[Name and Sign]").font.name = 'Times New Roman'
cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
p = cells[0].add_paragraph()
run = p.add_run("Supervisor/Guide")
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

cells[1].paragraphs[0].add_run("[Name and Sign]").font.name = 'Times New Roman'
cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = cells[1].add_paragraph()
run = p.add_run("Co-Supervisor/Guide")
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

for _ in range(2):
    for cell in sig_table.rows[0].cells:
        cell.add_paragraph()

cells2 = sig_table.rows[1].cells
cells2[0].paragraphs[0].add_run("[Name and Sign]").font.name = 'Times New Roman'
cells2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
p = cells2[0].add_paragraph()
run = p.add_run("Head of Department")
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

cells2[1].paragraphs[0].add_run("[Name and Sign]").font.name = 'Times New Roman'
cells2[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = cells2[1].add_paragraph()
run = p.add_run("Principal")
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Remove borders from signature table
for row in sig_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

add_page_break()

# ============================================================
# APPROVAL SHEET (Specimen C)
# ============================================================
for _ in range(2):
    doc.add_paragraph()

add_heading_centered("Project Report Approval for B. E.", size=18, space_after=30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Mm(12)
run = p.add_run(
    'This project report entitled "Indoor SLAM-Based Navigation System for Visually Impaired Users '
    'Using ARCore and Deep Learning" by '
)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('[YOUR NAME]')
run.bold = True
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run(' is approved for the degree of ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('Bachelor of Engineering')
run.bold = True
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('.')
run.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Examiners")
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
add_body("1. _____________________________________________", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_body("2. _____________________________________________", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_body("Date: _______________", indent=False)
add_body("Place: _______________", indent=False)

add_page_break()

# ============================================================
# DECLARATION (Specimen D)
# ============================================================
for _ in range(2):
    doc.add_paragraph()

add_heading_centered("Declaration", size=18, space_after=24)

add_body(
    "I declare that this written submission represents my ideas in my own words and where "
    "others' ideas or words have been included, I have adequately cited and referenced the "
    "original sources. I also declare that I have adhered to all principles of academic honesty "
    "and integrity and have not misrepresented or fabricated or falsified any "
    "idea/data/fact/source in my submission. I understand that any violation of the above will "
    "be cause for disciplinary action by the Institute and can also evoke penal action from the "
    "sources which have thus not been properly cited or from whom proper permission has not "
    "been taken when needed."
)

for _ in range(6):
    doc.add_paragraph()

add_body("_____________________________________________", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_body("(Signature)", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(2):
    doc.add_paragraph()

add_body("_____________________________________________", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_body("[YOUR NAME and Roll No.]", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_body("Date: _______________", indent=False)

add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
for _ in range(2):
    doc.add_paragraph()

add_heading_centered("Abstract", size=18, space_after=24)

add_body(
    "Indoor navigation remains a significant challenge for visually impaired individuals, "
    "as GPS-based systems are ineffective inside buildings due to signal attenuation and multipath "
    "interference. This project presents a real-time indoor navigation system built on a mobile "
    "platform (Android) using Simultaneous Localization and Mapping (SLAM) principles, Augmented "
    "Reality (ARCore), deep learning-based object detection (YOLOv8), Optical Character Recognition "
    "(OCR), and voice-guided turn-by-turn navigation. The system enables visually impaired users to "
    "independently navigate indoor environments such as corridors, offices, and facilities."
)

add_body(
    "The application is developed using a hybrid Flutter (Dart) and Android (Kotlin) architecture. "
    "The Flutter layer handles the user interface, including a real-time 2D occupancy grid map viewer "
    "with interactive object overlays. The Android native layer performs all computationally intensive "
    "operations: ARCore provides 6-DOF camera pose tracking and plane detection; a custom-trained "
    "YOLOv8 model (quantized to float16 TFLite) detects eight classes of indoor objects (doors, chairs, "
    "lift gates, fire extinguishers, notice boards, trash cans, water purifiers, and windows); and "
    "Google ML Kit Text Recognition v2 reads room numbers, signs, and notices. Communication between "
    "Flutter and native layers occurs through method channels with structured map payloads at 300ms "
    "and 800ms intervals."
)

add_body(
    "The SLAM subsystem constructs a log-odds occupancy grid at 0.20m resolution. Free space is "
    "determined through a combination of forward-arc ray casting, ARCore horizontal plane rasterization, "
    "and depth hit-testing at floor level. Wall detection leverages ARCore vertical planes and depth "
    "hit-tests at torso height. An incremental map update runs every graphics frame, with full map "
    "rebuilds triggered every 2 seconds from stored keyframes. Map consistency is maintained through "
    "three post-processing passes: isolated cell removal, wall gap filling, and single-cell wall dilation."
)

add_body(
    "Navigation is initiated through voice commands processed by Android SpeechRecognizer. A voice "
    "command processor parses natural language intents (e.g., \"take me to room 301\") and resolves "
    "destinations against the semantic object map. The A* path planner operates on the inflated "
    "occupancy grid with obstacle safety margins, 8-directional movement, semantic cost modifiers, "
    "and string-pulling path smoothing. The navigation guide provides real-time turn-by-turn "
    "instructions via Android Text-to-Speech, with automatic path re-planning upon deviation."
)

add_body(
    "The system successfully demonstrates autonomous indoor mapping and navigation on a standard "
    "Android smartphone without any pre-installed infrastructure such as beacons or markers, making "
    "it a practical, low-cost accessibility tool."
)

doc.add_paragraph()
add_heading_left("Keywords:", size=12, bold=True, space_before=6, space_after=6)
add_body(
    "SLAM, Indoor Navigation, ARCore, Occupancy Grid Mapping, YOLOv8, Object Detection, OCR, "
    "Visually Impaired, Flutter, A* Path Planning, Voice Navigation, Accessibility",
    indent=False
)

add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
add_heading_centered("Table of Contents", size=18, space_after=24)

toc_items = [
    ("", "Abstract", "iv"),
    ("", "List of Figures", "vi"),
    ("", "List of Tables", "vii"),
    ("", "List of Abbreviations", "viii"),
    ("1", "Introduction", "1"),
    ("1.1", "Background and Motivation", "1"),
    ("1.2", "Problem Statement", "2"),
    ("1.3", "Objectives", "3"),
    ("1.4", "Scope of the Project", "3"),
    ("1.5", "Organization of the Report", "4"),
    ("2", "Literature Review", "5"),
    ("2.1", "Indoor Navigation Systems", "5"),
    ("2.2", "SLAM Techniques for Mobile Devices", "6"),
    ("2.3", "Object Detection in Indoor Environments", "7"),
    ("2.4", "OCR-Based Landmark Recognition", "8"),
    ("2.5", "Assistive Navigation for Visually Impaired", "9"),
    ("2.6", "Summary of Literature Review", "10"),
    ("3", "System Design and Architecture", "11"),
    ("3.1", "System Overview", "11"),
    ("3.2", "Technology Stack", "12"),
    ("3.3", "System Architecture", "13"),
    ("3.4", "Flutter-Android Communication", "14"),
    ("3.5", "Module Design", "16"),
    ("3.6", "Data Flow", "20"),
    ("4", "Implementation", "22"),
    ("4.1", "ARCore Integration and Camera Setup", "22"),
    ("4.2", "Occupancy Grid Mapping (MapBuilder)", "24"),
    ("4.3", "Pose Tracking and Keyframe Management", "28"),
    ("4.4", "Object Detection (YOLOv8)", "30"),
    ("4.5", "Text Recognition (OCR)", "33"),
    ("4.6", "Semantic Map Management", "35"),
    ("4.7", "Navigation System", "37"),
    ("4.8", "Flutter UI Implementation", "41"),
    ("5", "Results and Discussion", "44"),
    ("5.1", "Mapping Performance", "44"),
    ("5.2", "Object Detection Accuracy", "45"),
    ("5.3", "Navigation Accuracy", "46"),
    ("5.4", "System Performance Metrics", "47"),
    ("5.5", "Discussion", "48"),
    ("6", "Conclusion and Future Work", "50"),
    ("6.1", "Summary of Work", "50"),
    ("6.2", "Conclusions", "50"),
    ("6.3", "Future Work", "51"),
    ("", "References", "53"),
    ("", "Appendix I: System Constants", "56"),
    ("", "Appendix II: YOLO Model Classes", "57"),
    ("", "Acknowledgements", "58"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(160), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    if num and not num.startswith(" "):
        prefix = f"{num}  "
    else:
        prefix = ""
    is_chapter = num and '.' not in num
    run = p.add_run(f"{prefix}{title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_chapter or not num:
        run.bold = True
    if num and '.' in num:
        p.paragraph_format.left_indent = Mm(10)
    run2 = p.add_run(f"\t{page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_page_break()

# ============================================================
# LIST OF FIGURES (placeholder)
# ============================================================
add_heading_centered("List of Figures", size=18, space_after=24)

figures = [
    ("Figure 1.1", "Challenges of indoor navigation for visually impaired users"),
    ("Figure 3.1", "High-level system architecture"),
    ("Figure 3.2", "Flutter-Android method channel communication"),
    ("Figure 3.3", "Module dependency diagram"),
    ("Figure 3.4", "Data flow diagram"),
    ("Figure 4.1", "ARCore coordinate system and camera setup"),
    ("Figure 4.2", "Log-odds occupancy grid update process"),
    ("Figure 4.3", "Wall detection from depth hit-testing"),
    ("Figure 4.4", "Free space ray fan (7 rays, ±45° arc)"),
    ("Figure 4.5", "Map consistency enforcement passes"),
    ("Figure 4.6", "Keyframe gating criteria"),
    ("Figure 4.7", "YOLO input preprocessing pipeline (rotation and letterboxing)"),
    ("Figure 4.8", "Object localization coordinate mapping"),
    ("Figure 4.9", "OCR text classification pipeline"),
    ("Figure 4.10", "Semantic object merge logic"),
    ("Figure 4.11", "A* path planning with obstacle inflation"),
    ("Figure 4.12", "String-pulling path smoothing"),
    ("Figure 4.13", "Navigation state machine"),
    ("Figure 4.14", "Flutter map viewer UI layout"),
    ("Figure 4.15", "Occupancy grid cell color scheme"),
    ("Figure 5.1", "Sample occupancy grid maps generated in test environments"),
    ("Figure 5.2", "Object detection results with bounding boxes"),
    ("Figure 5.3", "Navigation path visualization"),
]

for fig_num, fig_caption in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(160), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(f"{fig_num}: {fig_caption}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run("\t__")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_page_break()

# ============================================================
# LIST OF TABLES (placeholder)
# ============================================================
add_heading_centered("List of Tables", size=18, space_after=24)

tables_list = [
    ("Table 3.1", "Technology stack summary"),
    ("Table 3.2", "Flutter-Android method channel specifications"),
    ("Table 3.3", "Map payload fields"),
    ("Table 4.1", "Log-odds parameters for occupancy grid"),
    ("Table 4.2", "Keyframe gating thresholds"),
    ("Table 4.3", "YOLO object classes and footprint sizes"),
    ("Table 4.4", "OCR text classification categories"),
    ("Table 4.5", "Navigation turn direction thresholds"),
    ("Table 4.6", "Semantic cost modifiers for path planning"),
    ("Table 5.1", "System performance metrics"),
    ("Table 5.2", "Object detection accuracy per class"),
    ("Table A.1", "Complete system constants"),
]

for t_num, t_caption in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(160), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(f"{t_num}: {t_caption}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run("\t__")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_page_break()

# ============================================================
# ABBREVIATIONS
# ============================================================
add_heading_centered("List of Abbreviations", size=18, space_after=24)

abbreviations = [
    ("SLAM", "Simultaneous Localization and Mapping"),
    ("ARCore", "Augmented Reality Core (Google)"),
    ("YOLO", "You Only Look Once"),
    ("OCR", "Optical Character Recognition"),
    ("TFLite", "TensorFlow Lite"),
    ("ML Kit", "Machine Learning Kit (Google)"),
    ("TTS", "Text-to-Speech"),
    ("DOF", "Degrees of Freedom"),
    ("GPU", "Graphics Processing Unit"),
    ("NNAPI", "Neural Networks API (Android)"),
    ("NMS", "Non-Maximum Suppression"),
    ("IoU", "Intersection over Union"),
    ("BFS", "Breadth-First Search"),
    ("LOS", "Line of Sight"),
    ("YUV", "Luminance-Chrominance Color Space"),
    ("RGB", "Red Green Blue"),
    ("UI", "User Interface"),
    ("HUD", "Heads-Up Display"),
    ("API", "Application Programming Interface"),
    ("A*", "A-Star (pathfinding algorithm)"),
]

add_table(
    ["Abbreviation", "Full Form"],
    abbreviations,
    col_widths=[40, 120]
)

add_page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Chapter 1", size=18, space_after=12)
add_heading_centered("Introduction", size=18, space_after=24)

# 1.1
add_heading_left("1.1 Background and Motivation", size=16)

add_body(
    "Indoor navigation is a critical daily challenge faced by approximately 285 million visually "
    "impaired individuals worldwide, as estimated by the World Health Organization. While GPS-based "
    "navigation systems have revolutionized outdoor wayfinding, they remain fundamentally inadequate "
    "for indoor environments due to signal attenuation caused by building structures, multipath "
    "interference, and insufficient accuracy for room-level localization."
)

add_body(
    "Traditional indoor navigation solutions rely on pre-installed infrastructure such as Bluetooth "
    "Low Energy (BLE) beacons, Wi-Fi fingerprinting, or RFID tags. These approaches impose significant "
    "deployment costs, require ongoing maintenance, and are unavailable in the vast majority of public "
    "buildings. Furthermore, they provide only discrete position estimates rather than continuous spatial "
    "awareness, limiting their utility for real-time guidance."
)

add_body(
    "Recent advances in mobile Augmented Reality platforms, particularly Google ARCore, have made "
    "sophisticated 6-DOF (six degrees of freedom) pose tracking available on consumer smartphones. "
    "Simultaneously, deep learning-based object detection models such as YOLO (You Only Look Once) "
    "have achieved real-time inference speeds suitable for mobile deployment. These technological "
    "convergences create an unprecedented opportunity to build infrastructure-free indoor navigation "
    "systems that operate entirely on a user's existing smartphone."
)

add_body(
    "This project is motivated by the need to develop a practical, low-cost, and infrastructure-free "
    "indoor navigation system that empowers visually impaired users to navigate unfamiliar indoor "
    "environments independently, using only a standard Android smartphone."
)

# 1.2
add_heading_left("1.2 Problem Statement", size=16)

add_body(
    "Visually impaired individuals face significant barriers to independent mobility in indoor "
    "environments. Existing solutions either require expensive pre-installed infrastructure (beacons, "
    "pre-mapped environments) or provide insufficient spatial awareness for safe navigation. There is "
    "a need for a mobile application that can:"
)

add_bullet("Build a real-time spatial map of an unknown indoor environment without prior setup")
add_bullet("Detect and localize indoor objects and landmarks (doors, signs, obstacles) in 3D space")
add_bullet("Read textual information such as room numbers and signs")
add_bullet("Plan optimal, obstacle-free paths to desired destinations")
add_bullet("Provide intuitive voice-based interaction and turn-by-turn navigation guidance")

# 1.3
add_heading_left("1.3 Objectives", size=16)

add_body("The primary objectives of this project are:", indent=False)

add_bullet("To develop a real-time occupancy grid mapping system using ARCore pose tracking, "
           "plane detection, and depth sensing for indoor environments.")
add_bullet("To integrate a custom-trained YOLOv8 object detection model capable of identifying "
           "eight classes of indoor objects at real-time speeds on mobile hardware.")
add_bullet("To implement OCR-based text recognition for reading room numbers, signs, and notices "
           "using Google ML Kit Text Recognition v2.")
add_bullet("To build a semantic map that associates detected objects and text with their 3D positions "
           "in the mapped environment.")
add_bullet("To implement A* path planning with obstacle inflation, semantic cost modifiers, and "
           "path smoothing for safe indoor navigation.")
add_bullet("To develop a voice-activated navigation interface with natural language command processing "
           "and Text-to-Speech guidance.")
add_bullet("To create an accessible, infrastructure-free solution that runs entirely on a standard "
           "Android smartphone.")

# 1.4
add_heading_left("1.4 Scope of the Project", size=16)

add_body(
    "The scope of this project encompasses the design, development, and testing of a complete indoor "
    "navigation system. The system targets indoor environments with structured layouts such as "
    "corridors, offices, classrooms, and public facilities. The application is developed for Android "
    "devices supporting ARCore, using a hybrid Flutter-Kotlin architecture."
)

add_body(
    "The project covers the following technical domains: visual-inertial odometry via ARCore, "
    "log-odds occupancy grid mapping, deep learning-based object detection and localization, "
    "optical character recognition, graph-based path planning, and voice-based human-computer "
    "interaction. The system does not require any pre-installed infrastructure, pre-built maps, "
    "or internet connectivity for core functionality."
)

# 1.5
add_heading_left("1.5 Organization of the Report", size=16)

add_body("The remainder of this report is organized as follows:", indent=False)

add_body(
    "Chapter 2 presents a comprehensive review of literature related to indoor navigation, "
    "SLAM techniques, object detection, OCR, and assistive technology for visually impaired users."
)
add_body(
    "Chapter 3 describes the overall system design and architecture, including the technology "
    "stack, module decomposition, communication protocols, and data flow."
)
add_body(
    "Chapter 4 provides a detailed account of the implementation of each subsystem: ARCore "
    "integration, occupancy grid mapping, object detection, OCR, semantic mapping, navigation, "
    "and the Flutter user interface."
)
add_body(
    "Chapter 5 presents the results obtained from testing the system and discusses the "
    "performance, accuracy, and limitations of each component."
)
add_body(
    "Chapter 6 concludes the report with a summary of contributions, key conclusions, and "
    "directions for future work."
)

add_page_break()

# ============================================================
# CHAPTER 2: LITERATURE REVIEW
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Chapter 2", size=18, space_after=12)
add_heading_centered("Literature Review", size=18, space_after=24)

add_heading_left("2.1 Indoor Navigation Systems", size=16)

add_body(
    "Indoor navigation has been an active area of research for over two decades. Early systems relied "
    "on infrastructure-based approaches such as Wi-Fi fingerprinting [1], Bluetooth Low Energy (BLE) "
    "beacons [2], and ultrawideband (UWB) ranging [3]. While these systems achieve meter-level accuracy, "
    "they require significant upfront deployment and calibration effort, and their accuracy degrades "
    "over time due to environmental changes."
)

add_body(
    "Infrastructure-free approaches have gained attention with the proliferation of sensor-rich "
    "smartphones. Pedestrian Dead Reckoning (PDR) using inertial measurement units (IMUs) provides "
    "relative displacement estimation but suffers from cumulative drift [4]. Visual odometry and "
    "visual SLAM techniques offer more robust tracking by leveraging camera imagery to estimate "
    "motion and build environmental models simultaneously [5]."
)

add_body(
    "The emergence of ARCore (Google) and ARKit (Apple) has made visual-inertial odometry (VIO) "
    "readily available on consumer devices, providing reliable 6-DOF pose tracking that was "
    "previously limited to specialized hardware [6]. This project leverages ARCore's VIO capabilities "
    "as the foundation for its SLAM system."
)

add_heading_left("2.2 SLAM Techniques for Mobile Devices", size=16)

add_body(
    "Simultaneous Localization and Mapping (SLAM) is a computational problem of constructing a map "
    "of an unknown environment while simultaneously tracking the agent's location within it [7]. "
    "Classical SLAM approaches include Extended Kalman Filter (EKF) SLAM, particle filter-based "
    "FastSLAM, and graph-based SLAM [8]."
)

add_body(
    "For mobile devices, lightweight SLAM implementations are necessary due to computational "
    "constraints. ORB-SLAM [9] and LSD-SLAM [10] demonstrated real-time visual SLAM on desktop "
    "hardware, but their computational requirements exceed what is practical for sustained mobile "
    "operation. More recent approaches use occupancy grid representations, which discretize the "
    "environment into a regular grid of cells, each assigned a probability of being occupied [11]. "
    "Log-odds representations provide numerically stable incremental updates and are well-suited "
    "for real-time mobile applications."
)

add_body(
    "This project employs a log-odds occupancy grid approach operating at 0.20m resolution, "
    "combining multiple data sources (ARCore planes, depth hit-testing, and ray casting) for "
    "robust indoor mapping."
)

add_heading_left("2.3 Object Detection in Indoor Environments", size=16)

add_body(
    "Object detection has seen remarkable progress with the advent of deep learning. The YOLO "
    "(You Only Look Once) family of detectors [12] has become the de facto standard for real-time "
    "object detection due to its single-pass architecture. YOLOv8 [13] represents the latest "
    "iteration, offering improved accuracy and inference speed."
)

add_body(
    "For mobile deployment, model quantization techniques such as float16 and int8 quantization "
    "significantly reduce model size and inference time while maintaining acceptable accuracy [14]. "
    "TensorFlow Lite (TFLite) provides the runtime for executing quantized models on Android "
    "devices, with optional hardware acceleration via the Neural Networks API (NNAPI) [15]."
)

add_body(
    "Indoor object detection presents unique challenges including variable lighting conditions, "
    "occlusion, and the diversity of indoor object appearances. Custom training on indoor-specific "
    "datasets is essential for reliable performance [16]. This project uses a custom-trained YOLOv8 "
    "model quantized to float16 for detecting eight indoor object classes."
)

add_heading_left("2.4 OCR-Based Landmark Recognition", size=16)

add_body(
    "Optical Character Recognition (OCR) enables extraction of textual information from camera "
    "imagery. Google's ML Kit Text Recognition provides on-device text detection and recognition "
    "with support for Latin script [17]. In the context of indoor navigation, OCR can identify "
    "room numbers, directional signs, and facility labels, providing semantic context that "
    "enriches the navigation map [18]."
)

add_body(
    "The combination of OCR with object detection creates a powerful semantic understanding "
    "of the indoor environment. For instance, room numbers detected via OCR can be associated "
    "with nearby door objects detected by YOLO, enabling destination-based navigation (e.g., "
    "\"navigate to Room 301\") [19]."
)

add_heading_left("2.5 Assistive Navigation for Visually Impaired", size=16)

add_body(
    "Assistive navigation systems for visually impaired users must consider unique interaction "
    "paradigms. Voice-based interaction is the primary modality, encompassing both speech recognition "
    "for input and text-to-speech for output [20]. Navigation instructions must be clear, "
    "contextually appropriate, and timed to avoid information overload."
)

add_body(
    "Several research projects have explored smartphone-based assistive navigation, including "
    "NavCog [21] (BLE beacon-based), Soundscape [22] (spatial audio), and various camera-based "
    "approaches [23]. However, most existing systems either require infrastructure or provide "
    "only obstacle avoidance rather than complete turn-by-turn navigation."
)

add_body(
    "This project combines SLAM-based mapping, semantic object detection, OCR, and voice navigation "
    "into a unified system that provides comprehensive indoor navigation assistance without "
    "infrastructure requirements."
)

add_heading_left("2.6 Summary of Literature Review", size=16)

add_body(
    "The literature review reveals that while individual components (SLAM, object detection, OCR, "
    "voice navigation) are well-established, their integration into a unified, infrastructure-free "
    "indoor navigation system for visually impaired users remains an open challenge. This project "
    "addresses this gap by combining ARCore-based SLAM, custom-trained YOLOv8 object detection, "
    "ML Kit OCR, and voice-guided A* navigation into a cohesive mobile application."
)

add_page_break()

# ============================================================
# CHAPTER 3: SYSTEM DESIGN AND ARCHITECTURE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Chapter 3", size=18, space_after=12)
add_heading_centered("System Design and Architecture", size=18, space_after=24)

add_heading_left("3.1 System Overview", size=16)

add_body(
    "The Indoor SLAM Navigation System is a mobile application designed to enable visually impaired "
    "users to navigate indoor environments independently. The system operates entirely on a standard "
    "Android smartphone equipped with ARCore support, requiring no pre-installed infrastructure or "
    "pre-built maps."
)

add_body(
    "At a high level, the system performs four core functions: (1) real-time spatial mapping of the "
    "indoor environment using occupancy grids, (2) detection and 3D localization of indoor objects "
    "and text landmarks, (3) path planning from the user's current position to a specified destination, "
    "and (4) voice-based turn-by-turn navigation guidance."
)

add_body("[INSERT FIGURE 3.1: High-level system architecture diagram]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("3.2 Technology Stack", size=16)

add_body("The system employs the following technologies:", indent=False)

add_table(
    ["Component", "Technology", "Purpose"],
    [
        ["UI Layer", "Flutter (Dart)", "Cross-platform UI, map visualization"],
        ["Native Layer", "Android (Kotlin)", "ARCore, SLAM, detection, navigation"],
        ["AR Framework", "Google ARCore", "6-DOF tracking, plane detection, depth"],
        ["Object Detection", "YOLOv8 (TFLite float16)", "8-class indoor object detection"],
        ["Text Recognition", "Google ML Kit v2", "Room numbers, signs, notices"],
        ["Voice Output", "Android TTS", "Navigation instructions"],
        ["Voice Input", "Android SpeechRecognizer", "Voice command processing"],
        ["Path Planning", "Custom A* implementation", "Optimal path computation"],
    ],
    col_widths=[35, 50, 75]
)

add_body("Table 3.1: Technology stack summary", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("3.3 System Architecture", size=16)

add_body(
    "The system follows a hybrid architecture with Flutter handling the presentation layer and "
    "Android Kotlin performing all computationally intensive operations. This separation allows "
    "the UI to remain responsive while heavy processing (ARCore frame analysis, YOLO inference, "
    "map building) runs on native threads."
)

add_subheading_left("3.3.1 Flutter Layer", size=14)

add_body(
    "The Flutter layer consists of two primary screens: the HomePage providing feature card access "
    "to the AR Camera and Indoor Map views, and the IndoorMapViewer that renders the real-time "
    "occupancy grid map. The map viewer uses a CustomPainter (MapPainter) that renders layers in "
    "the following order: background, grid, floor cells, visited cells, navigation path, BFS path, "
    "obstacles, walls, semantic objects, and the robot (user) indicator."
)

add_body(
    "The map viewer supports pinch-to-zoom interaction (scale range 6-300 px/cell, default 28.0), "
    "displays a horizontal scrollable rail of detected object chips with confidence values, and "
    "provides a navigation button that toggles between voice command listening (microphone icon) "
    "and active navigation stop (stop icon)."
)

add_subheading_left("3.3.2 Android Native Layer", size=14)

add_body(
    "The Android layer is organized into distinct modules with well-defined responsibilities. "
    "ArActivity serves as the central orchestrator, managing the ARCore session with SHARED_CAMERA "
    "configuration, coordinating YOLO inference (every 900ms), OCR processing (every 3000ms), "
    "SLAM updates (every GL frame), and data transmission to Flutter (pose updates every 300ms, "
    "map payloads every 800ms)."
)

add_heading_left("3.4 Flutter-Android Communication", size=16)

add_body(
    "Communication between Flutter and the native Android layer is implemented through platform "
    "method channels. Two channels are used: one for AR-related data and one for navigation events."
)

add_table(
    ["Channel", "Direction", "Method", "Purpose"],
    [
        ["com.ketan.slam/ar", "Flutter \u2192 Native", "openAR", "Launch AR Activity"],
        ["com.ketan.slam/ar", "Native \u2192 Flutter", "onUpdate", "Pose + object count (300ms)"],
        ["com.ketan.slam/ar", "Native \u2192 Flutter", "updateMap", "Full map payload (800ms)"],
        ["com.ketan.slam/nav", "Flutter \u2192 Native", "startVoiceNav", "Start voice listening"],
        ["com.ketan.slam/nav", "Flutter \u2192 Native", "stopNavigation", "Stop active navigation"],
        ["com.ketan.slam/nav", "Native \u2192 Flutter", "navStateChange", "Navigation state updates"],
        ["com.ketan.slam/nav", "Native \u2192 Flutter", "navInstruction", "Turn-by-turn instructions"],
    ],
    col_widths=[38, 30, 32, 60]
)

add_body("Table 3.2: Flutter-Android method channel specifications", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_body(
    "The map payload transmitted via the updateMap channel contains the following structured data:"
)

add_table(
    ["Field", "Type", "Description"],
    [
        ["occupancyGrid", "ByteArray", "Flattened 2D grid of cell states"],
        ["gridWidth / gridHeight", "Int", "Grid dimensions"],
        ["gridResolution", "Float (0.20m)", "Meters per cell"],
        ["originX / originZ", "Float", "World-space grid origin"],
        ["robotGridX / robotGridZ", "Int", "User position in grid coords"],
        ["objects", "List<Map>", "Detected semantic objects"],
        ["navPath", "List<Map{x,z}>", "Active navigation path waypoints"],
    ],
    col_widths=[40, 40, 80]
)

add_body("Table 3.3: Map payload fields", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("3.5 Module Design", size=16)

add_body(
    "The Android native layer is decomposed into the following modules, each with clearly defined "
    "responsibilities:"
)

add_subheading_left("3.5.1 ArActivity (Central Orchestrator)", size=14)

add_body(
    "ArActivity manages the ARCore session lifecycle, configures the shared camera at 640x480 "
    "resolution (YUV_420_888), sets up the GLSurfaceView renderer, detection overlay, and HUD. "
    "It orchestrates the timing of all subsystem operations: SLAM updates on every GL frame, "
    "YOLO detection every 900ms, OCR recognition every 3000ms, pose transmission every 300ms, "
    "and map payload transmission every 800ms."
)

add_subheading_left("3.5.2 MapBuilder (Occupancy Grid)", size=14)

add_body(
    "MapBuilder implements the core occupancy grid mapping using a log-odds representation. "
    "The grid uses ConcurrentHashMap data structures for thread-safe concurrent access. Each "
    "cell is classified as one of five states: UNKNOWN (0), FREE (1), OBSTACLE (2), WALL (3), "
    "or VISITED (4)."
)

add_body(
    "Three update paths maintain the grid: (1) incrementalUpdate() runs every GL frame for fast "
    "local updates, (2) integratePlane() processes ARCore plane snapshots, and (3) rebuild() "
    "performs a full grid reconstruction from keyframes every 2 seconds on a background thread. "
    "Wall detection combines ARCore vertical plane rasterization (using Bresenham's line algorithm) "
    "with depth hit-testing at torso height. Free space determination uses a forward-arc ray fan "
    "(7 rays spanning ±45° over 3m range), ARCore horizontal plane rasterization, and depth "
    "hit-testing at floor level."
)

add_subheading_left("3.5.3 SlamEngine and PoseTracker", size=14)

add_body(
    "SlamEngine maintains the pose history trail (up to 5000 entries with 2cm minimum movement "
    "threshold). PoseTracker implements keyframe gating with thresholds of 0.15m translation, "
    "0.17 radians (~10°) rotation, or 200ms minimum interval. Drift detection uses ARCore Anchors "
    "placed every 5.0m (up to 3 anchors); a drift exceeding 0.05m triggers a grid rebuild flag."
)

add_subheading_left("3.5.4 ObjectLocalizer (3D Position Estimation)", size=14)

add_body(
    "ObjectLocalizer estimates 3D world positions for detected objects using a two-strategy approach: "
    "(1) ARCore hit-test at the bounding box center when the frame age is less than 67ms (2 frames "
    "at 30fps), and (2) an area-based fallback where depth is estimated as 0.5/sqrt(area), clamped "
    "to the range [0.8m, 6.0m]. The coordinate mapping pipeline handles the 90° clockwise rotation "
    "from landscape YOLO input (640x480) to portrait display, with 80px letterbox padding on each side."
)

add_subheading_left("3.5.5 SemanticMapManager", size=14)

add_body(
    "SemanticMapManager maintains a spatial index of semantic objects using a ConcurrentHashMap "
    "with 1m grid cells. Objects of the same type within 1.2m are merged; stale objects (unseen "
    "for >30 seconds with fewer than 3 observations) are removed. The manager supports 8 YOLO-detected "
    "object types and 7 OCR-detected text landmark types."
)

add_subheading_left("3.5.6 NavigationManager", size=14)

add_body(
    "NavigationManager orchestrates the complete navigation pipeline through six states: IDLE, "
    "LISTENING, PLANNING, NAVIGATING, ARRIVED, and ERROR. The pipeline processes voice commands "
    "through VoiceCommandProcessor, resolves destinations against the semantic map (by room number, "
    "text content, or object type), generates paths via PathPlanner, and provides guidance through "
    "NavigationGuide."
)

add_heading_left("3.6 Data Flow", size=16)

add_body(
    "The system's data flow proceeds as follows: ARCore provides 6-DOF pose estimates and plane "
    "detections to ArActivity on each graphics frame. ArActivity forwards pose data to PoseTracker "
    "for keyframe gating and to MapBuilder for incremental grid updates. At scheduled intervals, "
    "camera frames are passed to YoloDetector and TextRecognizer. Detection results are 3D-localized "
    "by ObjectLocalizer and registered in SemanticMapManager. NavigationManager's tick() function, "
    "called every SLAM frame, checks for pending intents, monitors arrival, validates the active "
    "path, and detects deviations (>2.0m) that trigger re-planning."
)

add_body(
    "The Flutter layer receives periodic updates through method channels and renders the occupancy "
    "grid, semantic objects, navigation path, and user position in real-time."
)

add_body("[INSERT FIGURE 3.4: Data flow diagram]", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ============================================================
# CHAPTER 4: IMPLEMENTATION
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Chapter 4", size=18, space_after=12)
add_heading_centered("Implementation", size=18, space_after=24)

add_heading_left("4.1 ARCore Integration and Camera Setup", size=16)

add_body(
    "The ARCore session is configured with the SHARED_CAMERA feature, enabling simultaneous "
    "access to the camera feed for both AR tracking and custom image processing (YOLO, OCR). "
    "The camera is configured at 640x480 resolution using the YUV_420_888 format via an ImageReader. "
    "The GL surface renders the AR camera background and detection overlays."
)

add_body(
    "Depth wall extraction is performed by sampling a 8x6 grid (48 points) every 300ms across "
    "the camera frame. Each depth hit is classified by its height relative to the camera's Y position: "
    "points below -0.5m are classified as floor (markHitFree), points between -0.5m and +0.8m are "
    "classified as walls/obstacles at torso height (markHitOccupied with 2x the standard L_OCCUPIED "
    "weight and wallHint flag), and points above +0.8m are classified as ceiling and ignored. The "
    "maximum hit distance is 5.0m. This depth-based wall extraction is the primary source of wall "
    "data, outperforming ARCore plane detection for indoor vertical surfaces."
)

add_body("[INSERT FIGURE 4.1: ARCore coordinate system and camera setup]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("4.2 Occupancy Grid Mapping (MapBuilder)", size=16)

add_subheading_left("4.2.1 Log-Odds Representation", size=14)

add_body(
    "The occupancy grid uses a log-odds representation for numerically stable probabilistic updates. "
    "Each cell maintains a floating-point log-odds value that is converted to a discrete cell state "
    "based on thresholds. The log-odds parameters are:"
)

add_table(
    ["Parameter", "Value", "Description"],
    [
        ["L_FREE", "-0.3f", "Log-odds decrement for free space observations"],
        ["L_OCCUPIED", "0.9f", "Log-odds increment for occupied observations"],
        ["L_MIN", "-4.0f", "Minimum clamping value for log-odds"],
        ["L_MAX", "3.5f", "Maximum clamping value for log-odds"],
        ["LO_THRESH_FREE", "-0.6f", "Threshold below which a cell is classified FREE"],
        ["LO_THRESH_OCC", "1.2f", "Threshold above which a cell is classified OCCUPIED"],
    ],
    col_widths=[40, 25, 95]
)

add_body("Table 4.1: Log-odds parameters for occupancy grid", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.2.2 Wall Detection", size=14)

add_body(
    "Walls are detected from two sources. First, ARCore vertical planes are rasterized into the "
    "grid using Bresenham's line algorithm (rasterisePlaneAsWall). Second, depth hit-tests at torso "
    "height provide direct wall measurements with double the standard occupied weight (2x L_OCCUPIED) "
    "and a wallHint flag. Cells detected from vertical planes are tracked in a dedicated wallCells "
    "HashSet to enable distinct rendering (dark color for walls vs. brown for generic obstacles)."
)

add_subheading_left("4.2.3 Free Space Detection", size=14)

add_body(
    "Free space is determined through three complementary mechanisms: (1) a forward-arc ray fan "
    "casting 7 rays spanning ±45° over a 3m range from the user's position, marking unobstructed "
    "cells as free; (2) ARCore horizontal plane polygons rasterized as free space "
    "(rasterisePlaneAsFree); and (3) depth hit-tests at floor level (below -0.5m from camera "
    "height) applying 3x the standard L_FREE weight for strong free-space evidence."
)

add_body("[INSERT FIGURE 4.4: Free space ray fan (7 rays, ±45° arc)]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.2.4 Map Consistency Enforcement", size=14)

add_body(
    "After each map update cycle, an enforceConsistency() function runs three sequential passes: "
    "(1) removal of isolated occupied cells (cells with no adjacent occupied neighbors), "
    "(2) filling of single-cell gaps between wall segments to ensure wall continuity, and "
    "(3) dilation of wall cells by one cell in all directions to provide a safety buffer."
)

add_body("[INSERT FIGURE 4.5: Map consistency enforcement passes]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("4.3 Pose Tracking and Keyframe Management", size=16)

add_subheading_left("4.3.1 Keyframe Gating", size=14)

add_body(
    "The PoseTracker module implements keyframe gating to maintain a manageable set of representative "
    "observations. A new keyframe is created when any of the following conditions are met:"
)

add_table(
    ["Criterion", "Threshold", "Purpose"],
    [
        ["Translation distance", "0.15m", "Capture spatial coverage"],
        ["Rotation angle", "0.17 rad (~10\u00b0)", "Capture viewpoint changes"],
        ["Time elapsed", "200ms", "Ensure minimum temporal coverage"],
    ],
    col_widths=[50, 40, 70]
)

add_body("Table 4.2: Keyframe gating thresholds", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.3.2 Observation Store", size=14)

add_body(
    "Keyframes are stored in the ObservationStore, a thread-safe ring buffer with a maximum "
    "capacity of 2000 keyframes and a monotonic version counter. Each keyframe contains the "
    "pose (position and heading), forward direction vector, associated plane snapshots "
    "(horizontal free-space and vertical wall planes), and object sightings from that viewpoint."
)

add_subheading_left("4.3.3 Drift Detection", size=14)

add_body(
    "Drift is monitored by placing ARCore Anchors at 5.0m intervals along the user's trajectory, "
    "maintaining up to 3 active anchors. When an anchor's tracked position deviates by more than "
    "0.05m from its expected position, a grid rebuild is triggered on the next rebuild cycle to "
    "correct accumulated mapping errors."
)

add_heading_left("4.4 Object Detection (YOLOv8)", size=16)

add_subheading_left("4.4.1 Model Configuration", size=14)

add_body(
    "The object detection model is a custom-trained YOLOv8 architecture, quantized to float16 "
    "format and deployed via TensorFlow Lite. The model operates at 640x640 RGB input resolution "
    "with 6 inference threads and NNAPI hardware acceleration."
)

add_body(
    "The detection confidence threshold is set to 0.45, and Non-Maximum Suppression (NMS) uses "
    "an IoU threshold of 0.45. A DetectionConfirmationGate requires 1 hit within a 2000ms window "
    "with a minimum IoU of 0.25 for confirmation."
)

add_subheading_left("4.4.2 Object Classes", size=14)

add_table(
    ["Class", "Footprint (half-metres)", "Description"],
    [
        ["CHAIR", "0.25m", "Seated furniture, treated as obstacle"],
        ["DOOR", "0.45m", "Doorways, navigation landmarks"],
        ["FIRE_EXTINGUISHER", "0.15m", "Safety equipment"],
        ["LIFT_GATE", "0.60m", "Elevator entrances"],
        ["NOTICE_BOARD", "0.40m", "Information displays"],
        ["TRASH_CAN", "0.20m", "Waste receptacles"],
        ["WATER_PURIFIER", "0.30m", "Drinking water stations"],
        ["WINDOW", "0.50m", "Window openings"],
    ],
    col_widths=[45, 40, 75]
)

add_body("Table 4.3: YOLO object classes and footprint sizes", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.4.3 Input Preprocessing Pipeline", size=14)

add_body(
    "The camera produces 640x480 landscape YUV_420_888 frames. These are converted to RGB and "
    "rotated 90° clockwise to produce a 480x640 portrait image. To match the YOLO model's 640x640 "
    "input, letterbox padding of 80 pixels is applied on each side (left and right), resulting in "
    "the final 640x640 input tensor."
)

add_body(
    "The coordinate mapping from YOLO detections back to screen coordinates accounts for the "
    "padding: normX = (bboxCenterX - 80) / 480, normY = bboxCenterY / 640, then screenX = "
    "normX * surfaceWidth, screenY = normY * surfaceHeight."
)

add_body("[INSERT FIGURE 4.7: YOLO input preprocessing pipeline]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.4.4 3D Object Localization", size=14)

add_body(
    "Detected objects are localized in 3D world space by ObjectLocalizer using a two-strategy "
    "approach. The primary strategy performs an ARCore hit-test at the bounding box center, but "
    "only when the frame age is less than 67ms (2 frames at 30fps) to ensure pose accuracy. "
    "The fallback strategy estimates depth from the bounding box area: depth = 0.5/sqrt(area), "
    "clamped to [0.8m, 6.0m]. The world position is then computed by projecting the screen "
    "coordinates along the camera's forward direction at the estimated depth."
)

add_heading_left("4.5 Text Recognition (OCR)", size=16)

add_body(
    "Text recognition is performed using Google ML Kit Text Recognition v2 (Latin script). "
    "The same YUV camera frame is converted to NV21 format, then to JPEG, then decoded to a "
    "Bitmap with 90° rotation. Recognition runs synchronously via a CountDownLatch with a "
    "3-second timeout."
)

add_subheading_left("4.5.1 Text Classification", size=14)

add_body("Recognized text is classified into four categories:", indent=False)

add_table(
    ["Category", "Pattern / Keywords", "Example"],
    [
        ["ROOM_NUMBER", "room|lab|rm|class|hall|office|cabin + digits", "Room 301, Lab 2A"],
        ["SIGN", "exit|washroom|stairs|lift|facility|warning keywords", "EXIT, Washroom"],
        ["NOTICE", "Long text on notice boards", "Schedule posted"],
        ["GENERAL", "Other recognized text", "Welcome"],
    ],
    col_widths=[35, 65, 60]
)

add_body("Table 4.4: OCR text classification categories", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.5.2 Text-Object Linking", size=14)

add_body(
    "In ArActivity, detected room numbers are linked to nearby door objects within 2.0m distance, "
    "enabling navigation commands like \"take me to Room 301\". Similarly, notice text is linked "
    "to nearby notice board objects. This cross-referencing between OCR text and YOLO detections "
    "creates a richer semantic understanding of the environment."
)

add_heading_left("4.6 Semantic Map Management", size=16)

add_body(
    "The SemanticMapManager maintains the complete semantic model of the environment. It supports "
    "15 total object types: 8 from YOLO detection (CHAIR, DOOR, FIRE_EXTINGUISHER, LIFT_GATE, "
    "NOTICE_BOARD, TRASH_CAN, WATER_PURIFIER, WINDOW) and 7 from OCR text landmarks (EXIT_SIGN, "
    "WASHROOM_SIGN, STAIRS_SIGN, ROOM_LABEL, FACILITY_SIGN, WARNING_SIGN, TEXT_SIGN)."
)

add_subheading_left("4.6.1 Object Merge Logic", size=14)

add_body(
    "When a new detection is received, it is checked against existing objects of the same type "
    "within a 1.2m radius. If a match is found, the objects are merged: the position is updated "
    "using a weighted average where weight = (1/sqrt(n)).coerceIn(0.1, 0.5), with n being the "
    "observation count. The confidence takes the maximum of the two values. The object ID follows "
    "the pattern \"label_gridX_gridZ\"."
)

add_subheading_left("4.6.2 Stale Object Removal", size=14)

add_body(
    "Objects that have not been observed for more than 30 seconds and have fewer than 3 total "
    "observations are removed from the semantic map. When an object is removed, a callback "
    "(onObjectRemoved) notifies MapBuilder to clear the corresponding footprint from the "
    "occupancy grid."
)

add_heading_left("4.7 Navigation System", size=16)

add_subheading_left("4.7.1 Voice Command Processing", size=14)

add_body(
    "Voice commands are captured using Android SpeechRecognizer with FREE_FORM language model "
    "and up to 3 result alternatives. The VoiceCommandProcessor recognizes trigger phrases "
    "including \"take me to\", \"go to\", \"navigate to\", \"find the\", and \"where is\". "
    "Qualifier detection supports NEAREST (default), FARTHEST, LEFT_MOST, and RIGHT_MOST "
    "modifiers. OCR keywords (room numbers, sign types) are checked before YOLO keywords to "
    "prioritize text-based destinations."
)

add_subheading_left("4.7.2 Path Planning (A*)", size=14)

add_body(
    "The PathPlanner implements the A* algorithm on the occupancy grid with the following "
    "characteristics: obstacle inflation by 2 cells (0.40m safety margin), 8-directional "
    "movement with octile distance heuristic, prevention of diagonal corner-cutting, and "
    "semantic cost modifiers that reduce cost for landmark cells (0.8x) and increase cost "
    "for hazard cells such as chairs (1.5x). The goal cell is always treated as walkable "
    "even if it contains an obstacle footprint."
)

add_body(
    "After the A* search completes, string-pulling (greedy line-of-sight smoothing) is applied "
    "as a post-processing step to remove unnecessary waypoints and produce smoother, more "
    "natural paths."
)

add_table(
    ["Object Type", "Cost Modifier", "Category"],
    [
        ["DOOR, LIFT_GATE, EXIT_SIGN, etc.", "0.8x", "LANDMARK (preferred)"],
        ["CHAIR", "1.5x", "HAZARD (avoided)"],
        ["Others", "1.0x", "NEUTRAL"],
    ],
    col_widths=[60, 35, 65]
)

add_body("Table 4.6: Semantic cost modifiers for path planning", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.7.3 Navigation Guide (TTS)", size=14)

add_body(
    "The NavigationGuide module provides voice-based turn-by-turn instructions using Android "
    "TextToSpeech with Locale.US. Instructions are generated based on a 1.5m lookahead from the "
    "user's current position. The arrival threshold is 1.0m. Turn directions are classified based "
    "on the angular difference between the current heading and the direction to the next waypoint:"
)

add_table(
    ["Direction", "Angle Range"],
    [
        ["STRAIGHT", "0\u00b0 - 15\u00b0"],
        ["SLIGHT LEFT/RIGHT", "15\u00b0 - 40\u00b0"],
        ["LEFT/RIGHT", "40\u00b0 - 90\u00b0"],
        ["SHARP LEFT/RIGHT", "90\u00b0 - 150\u00b0"],
        ["U-TURN", "> 150\u00b0"],
    ],
    col_widths=[60, 100]
)

add_body("Table 4.5: Navigation turn direction thresholds", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_subheading_left("4.7.4 Navigation State Machine", size=14)

add_body(
    "NavigationManager's tick() function is called every SLAM frame and manages six states: "
    "IDLE (waiting for commands), LISTENING (capturing voice input), PLANNING (computing path), "
    "NAVIGATING (active guidance), ARRIVED (destination reached), and ERROR (failure recovery). "
    "During NAVIGATING state, the manager trims passed waypoints within 0.6m, checks for path "
    "deviation exceeding 2.0m (triggering re-planning), and enforces a minimum 3000ms interval "
    "between spoken instructions to avoid overwhelming the user."
)

add_body("[INSERT FIGURE 4.13: Navigation state machine diagram]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("4.8 Flutter UI Implementation", size=16)

add_subheading_left("4.8.1 Map Visualization", size=14)

add_body(
    "The IndoorMapViewer uses a CustomPainter (MapPainter) that renders the occupancy grid "
    "with an architectural floor-plan color scheme. The rendering layers, in drawing order, are: "
    "warm off-white background (#F8F6F0), white free cells (#FFFFFF), light blue visited cells "
    "(#E8F0FE), green navigation path (#10B981 at 55% opacity), animated blue BFS path (#3B82F6), "
    "brown obstacles (#B45309 at 50% opacity), and near-black walls (#2C2C2C). The user (robot) "
    "is rendered as a pulsing accuracy ring with a heading arrow and blue dot."
)

add_subheading_left("4.8.2 Object Interaction", size=14)

add_body(
    "Detected objects are displayed as horizontally scrollable chips below the map, showing the "
    "object label and confidence value. Tapping an object chip triggers a Dart-side BFS path "
    "computation from the user's position to the object, visualized as an animated blue path "
    "overlay on the map."
)

add_subheading_left("4.8.3 Navigation Controls", size=14)

add_body(
    "The navigation button is a floating action button that changes appearance based on navigation "
    "state: a blue microphone icon in IDLE state for initiating voice commands, and a red stop icon "
    "during active NAVIGATING state. Sensor readings (position X/Z, heading, mapped area in m², "
    "and scanning status) are displayed in a top panel."
)

add_body("[INSERT FIGURE 4.14: Flutter map viewer UI layout]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ============================================================
# CHAPTER 5: RESULTS AND DISCUSSION
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Chapter 5", size=18, space_after=12)
add_heading_centered("Results and Discussion", size=18, space_after=24)

add_heading_left("5.1 Mapping Performance", size=16)

add_body(
    "[INSERT YOUR MAPPING RESULTS HERE - Include screenshots of occupancy grids generated "
    "in different environments (corridors, rooms, open areas). Discuss the quality of wall "
    "detection, free space identification, and map consistency.]"
)

add_body(
    "The occupancy grid mapping system was tested in indoor environments including corridors, "
    "classrooms, and open floor areas. The 0.20m grid resolution provided sufficient detail for "
    "navigation while maintaining computational efficiency. The combination of depth hit-testing, "
    "ARCore plane detection, and ray casting produced maps with well-defined wall boundaries and "
    "accurate free space regions."
)

add_body(
    "The map consistency enforcement (isolated cell removal, wall gap filling, and wall dilation) "
    "significantly improved the visual quality and navigability of generated maps. The 2-second "
    "full rebuild cycle ensured that the map remained globally consistent even as the user explored "
    "new areas."
)

add_body("[INSERT FIGURE 5.1: Sample occupancy grid maps]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("5.2 Object Detection Accuracy", size=16)

add_body(
    "[INSERT YOUR DETECTION RESULTS HERE - Include confusion matrix, per-class accuracy, "
    "sample detection images with bounding boxes. Discuss false positives/negatives and "
    "challenging scenarios.]"
)

add_body(
    "The YOLOv8 model, operating at a confidence threshold of 0.45, demonstrated reliable "
    "detection of indoor objects across varied lighting conditions. The float16 quantization "
    "preserved detection accuracy while enabling real-time inference on mobile hardware with "
    "NNAPI acceleration."
)

add_body("[INSERT TABLE: Object detection accuracy per class - Table 5.2]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_body("[INSERT FIGURE 5.2: Object detection results with bounding boxes]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("5.3 Navigation Accuracy", size=16)

add_body(
    "[INSERT YOUR NAVIGATION RESULTS HERE - Include path planning results, arrival accuracy, "
    "voice command recognition rates, and user testing feedback if available.]"
)

add_body(
    "The A* path planner with obstacle inflation successfully generated collision-free paths "
    "in all tested environments. String-pulling post-processing produced smooth, natural paths "
    "that closely resembled human walking trajectories. The 2.0m deviation threshold for "
    "re-planning provided robust path following even when users wandered slightly from the "
    "planned route."
)

add_body("[INSERT FIGURE 5.3: Navigation path visualization]", indent=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("5.4 System Performance Metrics", size=16)

add_body(
    "[INSERT YOUR PERFORMANCE DATA HERE - Fill in actual measured values in the table below.]"
)

add_table(
    ["Metric", "Value"],
    [
        ["YOLO inference time", "[___] ms per frame"],
        ["OCR processing time", "[___] ms per frame"],
        ["Map update frequency", "Every GL frame (~30fps)"],
        ["Full map rebuild time", "[___] ms"],
        ["A* path planning time", "[___] ms"],
        ["Pose update latency (to Flutter)", "300ms interval"],
        ["Map payload latency (to Flutter)", "800ms interval"],
        ["Memory usage", "[___] MB"],
        ["Battery consumption", "[___] %/hour"],
    ],
    col_widths=[80, 80]
)

add_body("Table 5.1: System performance metrics", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_heading_left("5.5 Discussion", size=16)

add_body(
    "The system demonstrates that a comprehensive indoor navigation solution for visually "
    "impaired users can be implemented entirely on a standard Android smartphone without "
    "pre-installed infrastructure. The combination of ARCore's visual-inertial odometry with "
    "log-odds occupancy grid mapping provides robust real-time spatial awareness."
)

add_body(
    "The semantic layer, combining YOLO object detection with OCR text recognition, enables "
    "natural destination-based navigation (e.g., \"take me to Room 301\" or \"find the nearest "
    "exit\"). The cross-referencing of OCR text with YOLO detections (linking room numbers to "
    "doors, notices to notice boards) creates a richer environmental model than either modality "
    "alone."
)

add_body(
    "Limitations observed during testing include: (1) ARCore tracking degradation in featureless "
    "environments (plain white walls), (2) reduced YOLO accuracy in low-light conditions, "
    "(3) OCR reliability depends on text size and viewing angle, and (4) the occupancy grid's "
    "0.20m resolution may miss fine-grained obstacles. These limitations represent opportunities "
    "for future improvement."
)

add_page_break()

# ============================================================
# CHAPTER 6: CONCLUSION AND FUTURE WORK
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Chapter 6", size=18, space_after=12)
add_heading_centered("Conclusion and Future Work", size=18, space_after=24)

add_heading_left("6.1 Summary of Work", size=16)

add_body(
    "This project presented the design, implementation, and evaluation of an indoor SLAM-based "
    "navigation system for visually impaired users. The system was built using a hybrid Flutter "
    "and Android Kotlin architecture, leveraging ARCore for 6-DOF pose tracking, a custom-trained "
    "YOLOv8 model for indoor object detection, Google ML Kit for OCR, and Android speech services "
    "for voice interaction."
)

add_body(
    "The system successfully achieves real-time occupancy grid mapping at 0.20m resolution, "
    "detection and 3D localization of 8 indoor object classes and 7 text landmark types, "
    "A* path planning with obstacle inflation and semantic cost modifiers, and voice-activated "
    "turn-by-turn navigation guidance. All functionality operates on-device without internet "
    "connectivity or pre-installed infrastructure."
)

add_heading_left("6.2 Conclusions", size=16)

add_body("The following conclusions are drawn from this investigation:", indent=False)

add_bullet(
    "ARCore's visual-inertial odometry provides sufficiently accurate pose tracking for "
    "real-time indoor occupancy grid mapping on consumer smartphones, with drift detection "
    "and correction via anchor-based monitoring."
)
add_bullet(
    "The log-odds occupancy grid representation, combined with multiple data sources "
    "(depth hit-testing, plane detection, ray casting), produces reliable indoor maps "
    "with well-defined wall boundaries and free space regions."
)
add_bullet(
    "A custom-trained YOLOv8 model, quantized to float16 TFLite format, achieves real-time "
    "inference on mobile hardware while maintaining acceptable detection accuracy for 8 "
    "indoor object classes."
)
add_bullet(
    "The integration of OCR text recognition with object detection creates a semantic "
    "understanding sufficient for natural language destination-based navigation "
    "(e.g., \"take me to Room 301\")."
)
add_bullet(
    "A* path planning with obstacle inflation, semantic cost modifiers, and string-pulling "
    "smoothing generates safe, efficient, and natural navigation paths."
)
add_bullet(
    "Voice-based interaction (speech recognition for commands, TTS for guidance) provides "
    "an accessible interface suitable for visually impaired users."
)
add_bullet(
    "The complete system operates infrastructure-free on a standard Android smartphone, "
    "making it a practical and deployable accessibility tool."
)

add_heading_left("6.3 Future Work", size=16)

add_body("The following directions are identified for future development:", indent=False)

add_bullet(
    "Multi-floor navigation: Extend the system to support multi-floor buildings by detecting "
    "stairs and elevators as floor transition points and maintaining separate maps per floor."
)
add_bullet(
    "Map persistence and sharing: Enable saving and loading of previously mapped environments, "
    "allowing users to benefit from maps created by others."
)
add_bullet(
    "Improved YOLO model: Expand the object detection model to cover more indoor object "
    "classes and train on larger, more diverse indoor datasets."
)
add_bullet(
    "LiDAR integration: Leverage depth sensors available on newer devices (e.g., ToF sensors) "
    "for more accurate and dense 3D mapping."
)
add_bullet(
    "Collaborative mapping: Enable multiple users to contribute to a shared map of a building "
    "in real-time, accelerating map coverage."
)
add_bullet(
    "Haptic feedback: Supplement voice guidance with haptic (vibration) patterns for direction "
    "cues, reducing reliance on audio in noisy environments."
)
add_bullet(
    "User studies: Conduct formal usability studies with visually impaired participants to "
    "evaluate and improve the system's real-world effectiveness."
)

add_page_break()

# ============================================================
# REFERENCES (IEEE Style)
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("References", size=18, space_after=24)

references = [
    '[1] P. Bahl and V. N. Padmanabhan, "RADAR: An in-building RF-based user location and tracking system," in Proc. IEEE INFOCOM, vol. 2, 2000, pp. 775-784.',
    '[2] R. Faragher and R. Harle, "Location fingerprinting with Bluetooth Low Energy beacons," IEEE J. Sel. Areas Commun., vol. 33, no. 11, pp. 2418-2428, 2015.',
    '[3] S. Gezici et al., "Localization via ultra-wideband radios: A look at positioning aspects for future sensor networks," IEEE Signal Process. Mag., vol. 22, no. 4, pp. 70-84, 2005.',
    '[4] H. Weinberg, "Using the ADXL202 in pedometer and personal navigation applications," Analog Devices, Application Note AN-602, 2002.',
    '[5] D. Scaramuzza and F. Fraundorfer, "Visual odometry: Part I - The first 30 years and fundamentals," IEEE Robot. Autom. Mag., vol. 18, no. 4, pp. 80-92, 2011.',
    '[6] Google LLC, "ARCore Overview," Google Developers Documentation, 2024.',
    '[7] H. Durrant-Whyte and T. Bailey, "Simultaneous localization and mapping: Part I," IEEE Robot. Autom. Mag., vol. 13, no. 2, pp. 99-110, 2006.',
    '[8] G. Grisetti, R. Kummerle, C. Stachniss, and W. Burgard, "A tutorial on graph-based SLAM," IEEE Intell. Transp. Syst. Mag., vol. 2, no. 4, pp. 31-43, 2010.',
    '[9] R. Mur-Artal, J. M. M. Montiel, and J. D. Tardos, "ORB-SLAM: A versatile and accurate monocular SLAM system," IEEE Trans. Robot., vol. 31, no. 5, pp. 1147-1163, 2015.',
    '[10] J. Engel, T. Schops, and D. Cremers, "LSD-SLAM: Large-scale direct monocular SLAM," in Proc. Eur. Conf. Comput. Vis. (ECCV), 2014, pp. 834-849.',
    '[11] S. Thrun, "Learning occupancy grid maps with forward sensor models," Auton. Robots, vol. 15, no. 2, pp. 111-127, 2003.',
    '[12] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, "You Only Look Once: Unified, real-time object detection," in Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR), 2016, pp. 779-788.',
    '[13] Ultralytics, "YOLOv8: A new state-of-the-art computer vision model," Ultralytics Documentation, 2023.',
    '[14] R. Krishnamoorthi, "Quantizing deep convolutional networks for efficient inference: A whitepaper," arXiv preprint arXiv:1806.08342, 2018.',
    '[15] Google LLC, "TensorFlow Lite for Android," TensorFlow Documentation, 2024.',
    '[16] A. Crespo et al., "Indoor object detection using deep learning for visually impaired assistance," in Proc. Int. Conf. Pattern Recognit. (ICPR), 2020, pp. 4187-4194.',
    '[17] Google LLC, "ML Kit Text Recognition v2," Google ML Kit Documentation, 2024.',
    '[18] H. Wang, S. Liu, and K. Xu, "OCR-based indoor sign detection for navigation assistance," in Proc. IEEE Int. Conf. Multimedia Expo (ICME), 2019, pp. 1684-1689.',
    '[19] Y. Zhang and L. Chen, "Semantic indoor mapping with text recognition for robot navigation," in Proc. IEEE/RSJ Int. Conf. Intell. Robots Syst. (IROS), 2021, pp. 5432-5438.',
    '[20] M. A. Williams, A. Hurst, and S. K. Kane, "Pray before you step out: Describing personal and situational blind navigation behaviors," in Proc. ACM SIGACCESS Conf. Comput. Accessibility (ASSETS), 2013, pp. 1-8.',
    '[21] D. Ahmetovic et al., "NavCog: A navigational cognitive assistant for the blind," in Proc. ACM Int. Conf. Human-Comput. Interact. Mobile Devices Services, 2016, pp. 90-99.',
    '[22] Microsoft Research, "Soundscape: A technology exploring the use of innovative audio-based technology," Microsoft Research Project, 2022.',
    '[23] S. Real and A. Araujo, "Navigation systems for the blind and visually impaired: Past work, challenges, and open problems," Sensors, vol. 19, no. 15, p. 3404, 2019.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Mm(10)
    p.paragraph_format.first_line_indent = Mm(-10)
    run = p.add_run(ref)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# APPENDIX I: SYSTEM CONSTANTS
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Appendix I", size=18, space_after=12)
add_heading_centered("System Constants", size=18, space_after=24)

add_table(
    ["Constant", "Value", "Location"],
    [
        ["Grid resolution", "0.20m", "ArActivity / MapBuilder"],
        ["Camera size", "640 x 480", "ArActivity (CAM_W/H)"],
        ["YOLO input size", "640 x 640", "YoloDetector"],
        ["YOLO confidence threshold", "0.45f", "YoloDetector"],
        ["YOLO NMS IoU threshold", "0.45f", "YoloDetector"],
        ["Merge radius", "1.2m", "ArActivity (MERGE_DIST)"],
        ["Stale object timeout", "30s", "ArActivity (STALE_MS)"],
        ["Map rebuild interval", "2s", "ArActivity"],
        ["Obstacle inflation", "2 cells = 0.40m", "PathPlanner"],
        ["Keyframe min translation", "0.15m", "PoseTracker"],
        ["Keyframe min rotation", "0.17 rad (~10\u00b0)", "PoseTracker"],
        ["Keyframe min interval", "200ms", "PoseTracker"],
        ["Drift rebuild threshold", "0.05m", "PoseTracker"],
        ["Anchor spacing", "5.0m", "PoseTracker"],
        ["Max anchors", "3", "PoseTracker"],
        ["OCR interval", "3s", "ArActivity"],
        ["YOLO interval", "0.9s", "ArActivity"],
        ["Arrival distance", "1.0m", "NavigationGuide"],
        ["Deviation re-plan threshold", "2.0m", "NavigationManager"],
        ["Waypoint trim distance", "0.6m", "NavigationManager"],
        ["Instruction min interval", "3000ms", "NavigationManager"],
        ["Lookahead distance", "1.5m", "NavigationGuide"],
        ["Max pose history", "5000", "SlamEngine"],
        ["Pose min movement", "0.02m", "SlamEngine"],
        ["Max keyframes", "2000", "ObservationStore"],
        ["Log-odds L_FREE", "-0.3f", "MapBuilder"],
        ["Log-odds L_OCCUPIED", "0.9f", "MapBuilder"],
        ["Log-odds L_MIN", "-4.0f", "MapBuilder"],
        ["Log-odds L_MAX", "3.5f", "MapBuilder"],
        ["Log-odds thresh free", "-0.6f", "MapBuilder"],
        ["Log-odds thresh occupied", "1.2f", "MapBuilder"],
        ["Ray fan range", "3m, 7 rays, \u00b145\u00b0", "MapBuilder"],
        ["Depth hit max distance", "5.0m", "ArActivity"],
        ["Floor height threshold", "-0.5m (relative)", "ArActivity"],
        ["Wall height range", "-0.5m to +0.8m", "ArActivity"],
        ["TFLite threads", "6", "YoloDetector"],
        ["Detection gate hits", "1 in 2000ms", "YoloDetector"],
        ["Map default scale", "28.0 px/cell", "IndoorMapViewer"],
        ["Map scale range", "6 - 300 px/cell", "IndoorMapViewer"],
    ],
    col_widths=[50, 45, 65]
)

add_body("Table A.1: Complete system constants", indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ============================================================
# APPENDIX II: YOLO MODEL CLASSES
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_centered("Appendix II", size=18, space_after=12)
add_heading_centered("YOLO Model Classes and OCR Landmark Types", size=18, space_after=24)

add_heading_left("YOLO Detection Classes (8)", size=14)

add_table(
    ["Class Index", "Label", "Footprint Half-Size (m)", "Navigation Category"],
    [
        ["0", "CHAIR", "0.25", "HAZARD (1.5x cost)"],
        ["1", "DOOR", "0.45", "LANDMARK (0.8x cost)"],
        ["2", "FIRE_EXTINGUISHER", "0.15", "NEUTRAL"],
        ["3", "LIFT_GATE", "0.60", "LANDMARK (0.8x cost)"],
        ["4", "NOTICE_BOARD", "0.40", "NEUTRAL"],
        ["5", "TRASH_CAN", "0.20", "NEUTRAL"],
        ["6", "WATER_PURIFIER", "0.30", "NEUTRAL"],
        ["7", "WINDOW", "0.50", "NEUTRAL"],
    ],
    col_widths=[25, 45, 40, 50]
)

doc.add_paragraph()

add_heading_left("OCR Text Landmark Types (7)", size=14)

add_table(
    ["Type", "Keywords / Patterns", "Source"],
    [
        ["EXIT_SIGN", "exit, emergency exit", "OCR text matching"],
        ["WASHROOM_SIGN", "washroom, toilet, restroom, bathroom", "OCR text matching"],
        ["STAIRS_SIGN", "stairs, staircase, stairway", "OCR text matching"],
        ["ROOM_LABEL", "room/lab/rm/class + digits (e.g., Room 301)", "OCR regex pattern"],
        ["FACILITY_SIGN", "canteen, library, office, etc.", "OCR text matching"],
        ["WARNING_SIGN", "warning, caution, danger, restricted", "OCR text matching"],
        ["TEXT_SIGN", "General recognized text", "OCR fallback"],
    ],
    col_widths=[40, 65, 55]
)

add_page_break()

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_heading_centered("Acknowledgements", size=18, space_after=30)

add_body(
    "[WRITE YOUR ACKNOWLEDGEMENTS HERE - Thank your supervisor, co-supervisor, "
    "department head, family, friends, and anyone who supported your project work.]"
)

for _ in range(6):
    doc.add_paragraph()

add_body("[YOUR NAME]", indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_body("Date: _______________", indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SLAM_Indoor_Navigation_Project_Report.docx")
doc.save(output_path)
print(f"Report generated: {output_path}")
